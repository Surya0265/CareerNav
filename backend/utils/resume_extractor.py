import PyPDF2
import pdfplumber
from docx import Document
import os
import re


def _extract_section_entries(raw_text, section_keywords, stop_keywords):
    """Extract bullet-style entries for resume sections like experience or projects."""
    if not raw_text:
        return []

    lines = raw_text.splitlines()
    section_keywords_lower = [keyword.lower() for keyword in section_keywords]
    stop_keywords_lower = [keyword.lower() for keyword in stop_keywords]

    capturing = False
    buffer = []
    entries = []

    def flush_buffer():
        if buffer:
            combined = " ".join(buffer).strip()
            if combined:
                entries.append(combined)
            buffer.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if capturing:
                flush_buffer()
            continue

        lower_line = line.lower()

        # Detect start of section
        if any(lower_line.startswith(keyword) for keyword in section_keywords_lower):
            capturing = True
            flush_buffer()
            continue

        if capturing and any(lower_line.startswith(keyword) for keyword in stop_keywords_lower):
            flush_buffer()
            break

        if capturing:
            # Normalize bullet points
            if line[0] in {"-", "•", "*"}:
                flush_buffer()
                buffer.append(line.lstrip("-•* "))
            else:
                buffer.append(line)

    flush_buffer()

    # Remove duplicates while preserving order
    seen = set()
    deduped = []
    for entry in entries:
        if entry not in seen:
            deduped.append(entry)
            seen.add(entry)

    return deduped

def extract_text_from_pdf(file_path):
    """
    Extract text from PDF file using both PyPDF2 and pdfplumber for better coverage
    """
    text = ""
    
    try:
        # First try with pdfplumber (better for complex layouts)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # If pdfplumber didn't extract much text, try PyPDF2
        if len(text.strip()) < 100:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
    
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        return None
    
    return text.strip()

def extract_text_from_docx(file_path):
    """
    Extract text from DOCX file
    """
    try:
        doc = Document(file_path)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
    
    except Exception as e:
        print(f"Error extracting DOCX text: {str(e)}")
        return None

def extract_resume_text(file_path):
    """
    Main function to extract text from resume file
    Supports PDF and DOCX formats
    """
    if not os.path.exists(file_path):
        return None
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        print(f"Unsupported file format: {file_extension}")
        return None

def clean_extracted_text(text):
    """
    Clean and normalize extracted text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters that might cause issues
    text = re.sub(r'[^\w\s@.,()-]', '', text)
    
    # Normalize line breaks
    text = text.replace('\n', ' ').replace('\r', ' ')
    
    return text.strip()

def get_comprehensive_skills_database():
    """
    Comprehensive database of technical skills and technologies with aliases
    """
    return {
        'programming_languages': [
            {'name': 'Python', 'aliases': ['python', 'py']},
            {'name': 'JavaScript', 'aliases': ['javascript', 'js', 'ecmascript', 'es6', 'es2015', 'es2020']},
            {'name': 'TypeScript', 'aliases': ['typescript', 'ts']},
            {'name': 'Java', 'aliases': ['java']},
            {'name': 'C++', 'aliases': ['c++', 'cpp', 'cplusplus']},
            {'name': 'C#', 'aliases': ['c#', 'csharp', 'c sharp']},
            {'name': 'C', 'aliases': ['c programming', ' c ']},
            {'name': 'Go', 'aliases': ['go', 'golang']},
            {'name': 'Rust', 'aliases': ['rust']},
            {'name': 'PHP', 'aliases': ['php']},
            {'name': 'Ruby', 'aliases': ['ruby']},
            {'name': 'Swift', 'aliases': ['swift']},
            {'name': 'Kotlin', 'aliases': ['kotlin']},
            {'name': 'Scala', 'aliases': ['scala']},
            {'name': 'R', 'aliases': ['r programming', ' r ']},
            {'name': 'MATLAB', 'aliases': ['matlab']},
            {'name': 'Dart', 'aliases': ['dart']},
            {'name': 'Perl', 'aliases': ['perl']},
            {'name': 'Lua', 'aliases': ['lua']},
            {'name': 'Haskell', 'aliases': ['haskell']},
            {'name': 'Shell/Bash', 'aliases': ['bash', 'shell', 'zsh', 'powershell']},
            {'name': 'SQL', 'aliases': ['sql']},
        ],
        
        'web_frameworks': [
            {'name': 'React', 'aliases': ['react', 'reactjs', 'react.js']},
            {'name': 'Angular', 'aliases': ['angular', 'angularjs', 'angular.js']},
            {'name': 'Vue.js', 'aliases': ['vue', 'vue.js', 'vuejs']},
            {'name': 'Express.js', 'aliases': ['express', 'express.js', 'expressjs']},
            {'name': 'Django', 'aliases': ['django']},
            {'name': 'Flask', 'aliases': ['flask']},
            {'name': 'FastAPI', 'aliases': ['fastapi', 'fast api']},
            {'name': 'Spring Boot', 'aliases': ['spring boot', 'spring', 'springframework']},
            {'name': 'Laravel', 'aliases': ['laravel']},
            {'name': 'Ruby on Rails', 'aliases': ['rails', 'ruby on rails', 'ror']},
            {'name': 'ASP.NET', 'aliases': ['asp.net', 'aspnet', 'asp net']},
            {'name': 'Next.js', 'aliases': ['next', 'next.js', 'nextjs']},
            {'name': 'Nuxt.js', 'aliases': ['nuxt', 'nuxt.js', 'nuxtjs']},
            {'name': 'Svelte', 'aliases': ['svelte', 'sveltekit']},
            {'name': 'Node.js', 'aliases': ['node.js', 'nodejs', 'node']},
        ],
        
        'web_technologies': [
            {'name': 'HTML', 'aliases': ['html', 'html5']},
            {'name': 'CSS', 'aliases': ['css', 'css3']},
            {'name': 'Sass', 'aliases': ['sass', 'scss']},
            {'name': 'Bootstrap', 'aliases': ['bootstrap']},
            {'name': 'Tailwind CSS', 'aliases': ['tailwind', 'tailwindcss', 'tailwind css']},
            {'name': 'jQuery', 'aliases': ['jquery']},
            {'name': 'Webpack', 'aliases': ['webpack']},
            {'name': 'Vite', 'aliases': ['vite']},
            {'name': 'GraphQL', 'aliases': ['graphql']},
            {'name': 'REST API', 'aliases': ['rest', 'rest api', 'restful']},
        ],
        
        'databases': [
            {'name': 'MongoDB', 'aliases': ['mongodb', 'mongo']},
            {'name': 'PostgreSQL', 'aliases': ['postgresql', 'postgres', 'psql']},
            {'name': 'MySQL', 'aliases': ['mysql']},
            {'name': 'SQLite', 'aliases': ['sqlite']},
            {'name': 'Redis', 'aliases': ['redis']},
            {'name': 'Cassandra', 'aliases': ['cassandra']},
            {'name': 'DynamoDB', 'aliases': ['dynamodb']},
            {'name': 'Oracle', 'aliases': ['oracle', 'oracle db']},
            {'name': 'Microsoft SQL Server', 'aliases': ['sql server', 'mssql', 'microsoft sql']},
            {'name': 'Elasticsearch', 'aliases': ['elasticsearch', 'elastic search']},
            {'name': 'Firebase', 'aliases': ['firebase', 'firestore']},
        ],
        
        'cloud_platforms': [
            {'name': 'AWS', 'aliases': ['aws', 'amazon web services', 'ec2', 's3', 'lambda']},
            {'name': 'Google Cloud', 'aliases': ['gcp', 'google cloud', 'google cloud platform']},
            {'name': 'Microsoft Azure', 'aliases': ['azure', 'microsoft azure']},
            {'name': 'Digital Ocean', 'aliases': ['digitalocean', 'digital ocean']},
            {'name': 'Heroku', 'aliases': ['heroku']},
            {'name': 'Vercel', 'aliases': ['vercel']},
            {'name': 'Netlify', 'aliases': ['netlify']},
        ],
        
        'devops_tools': [
            {'name': 'Docker', 'aliases': ['docker', 'containerization']},
            {'name': 'Kubernetes', 'aliases': ['kubernetes', 'k8s']},
            {'name': 'Jenkins', 'aliases': ['jenkins']},
            {'name': 'Git', 'aliases': ['git', 'github', 'gitlab', 'bitbucket']},
            {'name': 'Terraform', 'aliases': ['terraform']},
            {'name': 'Ansible', 'aliases': ['ansible']},
            {'name': 'CI/CD', 'aliases': ['ci/cd', 'continuous integration', 'continuous deployment']},
            {'name': 'Nginx', 'aliases': ['nginx']},
            {'name': 'Apache', 'aliases': ['apache', 'apache http']},
        ],
        
        'mobile_development': [
            {'name': 'React Native', 'aliases': ['react native', 'react-native']},
            {'name': 'Flutter', 'aliases': ['flutter']},
            {'name': 'Xamarin', 'aliases': ['xamarin']},
            {'name': 'Ionic', 'aliases': ['ionic']},
            {'name': 'Android Development', 'aliases': ['android', 'android studio']},
            {'name': 'iOS Development', 'aliases': ['ios', 'xcode']},
        ],
        
        'data_science_ml': [
            {'name': 'TensorFlow', 'aliases': ['tensorflow', 'tf']},
            {'name': 'PyTorch', 'aliases': ['pytorch', 'torch']},
            {'name': 'Scikit-learn', 'aliases': ['scikit-learn', 'sklearn', 'scikit learn']},
            {'name': 'Pandas', 'aliases': ['pandas']},
            {'name': 'NumPy', 'aliases': ['numpy']},
            {'name': 'Matplotlib', 'aliases': ['matplotlib']},
            {'name': 'Seaborn', 'aliases': ['seaborn']},
            {'name': 'Keras', 'aliases': ['keras']},
            {'name': 'OpenCV', 'aliases': ['opencv', 'cv2']},
            {'name': 'Jupyter', 'aliases': ['jupyter', 'jupyter notebook']},
        ],
        
        'design_3d_tools': [
            {'name': 'Blender', 'aliases': ['blender']},
            {'name': 'Photoshop', 'aliases': ['photoshop', 'adobe photoshop']},
            {'name': 'Illustrator', 'aliases': ['illustrator', 'adobe illustrator']},
            {'name': 'Figma', 'aliases': ['figma']},
            {'name': 'Sketch', 'aliases': ['sketch']},
            {'name': 'Maya', 'aliases': ['maya', 'autodesk maya']},
            {'name': '3ds Max', 'aliases': ['3ds max', '3dsmax']},
            {'name': 'Unity', 'aliases': ['unity', 'unity3d']},
            {'name': 'Unreal Engine', 'aliases': ['unreal', 'unreal engine', 'ue4', 'ue5']},
        ],
        
        'testing_frameworks': [
            {'name': 'Jest', 'aliases': ['jest']},
            {'name': 'Mocha', 'aliases': ['mocha']},
            {'name': 'Cypress', 'aliases': ['cypress']},
            {'name': 'Selenium', 'aliases': ['selenium']},
            {'name': 'Pytest', 'aliases': ['pytest']},
            {'name': 'JUnit', 'aliases': ['junit']},
            {'name': 'Postman', 'aliases': ['postman']},
        ],
        
        'soft_skills': [
            {'name': 'Project Management', 'aliases': ['project management', 'pm']},
            {'name': 'Agile', 'aliases': ['agile', 'scrum', 'kanban']},
            {'name': 'Leadership', 'aliases': ['leadership', 'team lead', 'management']},
            {'name': 'Communication', 'aliases': ['communication', 'presentation']},
            {'name': 'Problem Solving', 'aliases': ['problem solving', 'analytical']},
        ]
    }

def extract_skills_from_text(text):
    """
    Extract and normalize skills from text using comprehensive database
    """
    if not text:
        return []
    
    skills_db = get_comprehensive_skills_database()
    found_skills = set()  # Use set to avoid duplicates
    text_lower = text.lower()
    
    # Search through all categories
    for category, skills in skills_db.items():
        for skill_info in skills:
            skill_name = skill_info['name']
            
            # Check if any alias matches
            for alias in skill_info['aliases']:
                # Use word boundaries for better matching
                pattern = r'\b' + re.escape(alias.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.add(skill_name)
                    break  # Found this skill, no need to check other aliases
    
    return sorted(list(found_skills))

def get_skills_summary(skills_list):
    """
    Categorize skills by type for better presentation
    """
    if not skills_list:
        return {}
    
    skills_db = get_comprehensive_skills_database()
    categorized = {}
    
    for skill in skills_list:
        # Find which category this skill belongs to
        for category, skills in skills_db.items():
            for skill_info in skills:
                if skill_info['name'] == skill:
                    if category not in categorized:
                        categorized[category] = []
                    categorized[category].append(skill)
                    break
    
    return categorized

def _normalize_section_line(line):
    # Remove bullet characters and extra symbols
    cleaned = re.sub(r'^[\s•\-*–—\u2022\u2023\u25AA\u25CF\u25E6]+', '', line).strip()
    return cleaned

def _extract_section_entries(raw_text, section_headers, stop_headers):
    if not raw_text:
        return []

    lines = raw_text.splitlines()
    capture = False
    entries = []
    buffer = []

    def flush_buffer():
        merged = ' '.join(buffer).strip()
        normalized = _normalize_section_line(merged)
        if normalized:
            entries.append(normalized)

    lower_stop_headers = [header.lower() for header in stop_headers]
    lower_section_headers = [header.lower() for header in section_headers]

    for line in lines:
        stripped = line.strip()
        lower_line = stripped.lower()

        if any(header in lower_line for header in lower_section_headers):
            if capture and buffer:
                flush_buffer()
                buffer = []
            capture = True
            continue

        if capture:
            if any(lower_line.startswith(stop) for stop in lower_stop_headers):
                break
            if not stripped:
                if buffer:
                    flush_buffer()
                    buffer = []
            else:
                buffer.append(stripped)

    if capture and buffer:
        flush_buffer()

    # Deduplicate while preserving order and limit to avoid excessively long arrays
    seen = set()
    unique_entries = []
    for entry in entries:
        lowered = entry.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        unique_entries.append(entry)
        if len(unique_entries) >= 12:
            break

    return unique_entries

def extract_basic_info(text, raw_text=None):
    """
    Extract basic information from resume text using enhanced skill extraction
    """
    info = {
        'email': None,
        'phone': None,
        'skills': [],
        'skills_summary': {},
        'experience_keywords': [],
        'education_keywords': [],
        'experience_entries': [],
        'project_entries': []
    }
    
    if not text:
        return info
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        info['email'] = emails[0]
    
    # Extract phone number (basic pattern)
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    if phones:
        info['phone'] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
    
    # Extract skills using comprehensive database
    skills = extract_skills_from_text(text)
    info['skills'] = skills
    info['skills_summary'] = get_skills_summary(skills)
    
    # Experience keywords
    exp_keywords = ['experience', 'worked', 'developed', 'managed', 'led', 'created', 'built']
    text_lower = text.lower()
    for keyword in exp_keywords:
        if keyword in text_lower:
            info['experience_keywords'].append(keyword)
    
    # Education keywords
    edu_keywords = ['bachelor', 'master', 'phd', 'degree', 'university', 'college', 'education']
    for keyword in edu_keywords:
        if keyword in text_lower:
            info['education_keywords'].append(keyword)
    
    raw_source = raw_text or text
    stop_headers = [
        'education', 'certifications', 'skills', 'projects', 'project experience',
        'about', 'summary', 'objective', 'achievements', 'publications'
    ]

    experience_headers = [
        'experience', 'work experience', 'professional experience', 'employment history'
    ]
    info['experience_entries'] = _extract_section_entries(raw_source, experience_headers, stop_headers)

    project_headers = [
        'projects', 'project experience', 'personal projects', 'notable projects', 'selected projects'
    ]
    info['project_entries'] = _extract_section_entries(raw_source, project_headers, stop_headers)

    return info
